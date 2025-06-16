from docx import Document
import datetime
import os

def generate_word_report(scan_data: dict, domain: str) -> str:
    reports_dir = "reports"
    os.makedirs(reports_dir, exist_ok=True)

    doc = Document()
    doc.add_heading("🛡️ MCyberRecon Security Report", 0)
    doc.add_paragraph(f"📌 Domain: {domain}")
    doc.add_paragraph(f"🕓 Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("This document contains the results of a vulnerability and recon scan.\n")

    for section, content in scan_data.items():
        doc.add_heading(f"🔹 {section}", level=1)
        if isinstance(content, dict):
            for key, value in content.items():
                doc.add_paragraph(f"{key}: {value}", style='List Bullet')
        elif isinstance(content, list):
            for item in content:
                doc.add_paragraph(f"- {item}", style='List Bullet')
        else:
            doc.add_paragraph(str(content))

    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    file_name = f"MCyberRecon_{domain}_{timestamp}.docx"
    file_path = os.path.join(reports_dir, file_name)
    doc.save(file_path)

    return file_path